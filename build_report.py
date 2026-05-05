from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # --- 1. Title Page ---
    # Rubric Section 1
    title = doc.add_heading('Integrated DRRM Business Intelligence Dashboard for Biñan City', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    course = doc.add_paragraph('Course: Business Intelligence and Data Visualization')
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    group = doc.add_paragraph('Group Members: Funelas, Rivera, Torculas')
    group.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    date = doc.add_paragraph('Submission Date: May 2025')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # --- 2. Introduction and BI context ---
    # Rubric Section 2
    doc.add_heading('2. Introduction and BI Context', level=1)
    doc.add_paragraph(
        "Our project is a Business Intelligence dashboard for the City of Biñan, Laguna. "
        "Biñan is a city that is very prone to flooding because it is located right next to Laguna de Bay. "
        "Because of this, the city needs a better way to monitor flood incidents and rescue resources."
    )
    doc.add_paragraph(
        "The target users for this dashboard are the DRRMO Officers and City Planners of Biñan. "
        "We classified this as an Operational and Tactical BI Dashboard. "
        "It is operational because it helps staff monitor the daily activities and incidents as they happen. "
        "It is tactical because it helps the leaders decide where to put their rescue teams and check if the evacuation centers are already full. "
        "The time horizon is mostly daily and periodic, focusing on current emergencies and short-term planning."
    )
    
    # --- 3. Data used in the system ---
    # Rubric Section 3
    doc.add_heading('3. Data Used in the System', level=1)
    doc.add_paragraph("We used four core datasets to build the dashboard. These files contain all the information needed for mapping and analysis.")
    
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Dataset'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Key Fields'
    hdr_cells[3].text = 'Used For'
    
    data1 = [
        ('barangays.csv', 'Barangay data and hazard levels', 'Barangay, Population, HazardLevel', 'Risk scoring and hazard mapping'),
        ('facilities.csv', 'List of city facilities', 'Type, Capacity, Occupants', 'Capacity analysis and shelters'),
        ('rescue_teams.csv', 'Rescue team details', 'Specialty, Status', 'Readiness and deployment check'),
        ('incidents.csv', 'Historical incident logs', 'DateTime, Type, Severity', 'Trend analysis and impact logs')
    ]
    for row in data1:
        row_cells = table1.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        
    # --- 4. Data dictionary, dimensions, and measures ---
    # Rubric Section 4
    doc.add_heading('4. Data Dictionary, Dimensions, and Measures', level=1)
    doc.add_paragraph("Here are the important fields we used to slice and measure the data in our system:")
    
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Field'
    hdr_cells2[1].text = 'Data Type'
    hdr_cells2[2].text = 'Classification'
    hdr_cells2[3].text = 'Meaning'
    hdr_cells2[4].text = 'Used In'
    
    data2 = [
        ('Month', 'Text', 'Dimension', 'Month of the event', 'Trend chart filter'),
        ('Barangay', 'Text', 'Dimension', 'Location in Biñan', 'Mapping and ranking'),
        ('Severity', 'Text', 'Dimension', 'Danger level', 'Severity pie chart'),
        ('NumAffected', 'Integer', 'Measure', 'Number of people', 'Affected population KPI'),
        ('Capacity', 'Integer', 'Measure', 'Max shelter space', 'Gap analysis chart'),
        ('RiskScore', 'Float', 'Measure', 'Computed risk level', 'Prioritizing barangays')
    ]
    for row in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]
 
    # --- 5. Data processing and ETL ---
    # Rubric Section 5
    doc.add_heading('5. Data Processing and ETL', level=1)
    doc.add_paragraph(
        "For our data processing, we followed the BI concept of ETL (Extract, Transform, Load). "
        "In the Extraction phase, we used the Pandas library to load all four CSV files into our Jupyter Notebook. "
    )
    doc.add_paragraph(
        "During the cleaning and transformation phase, we found a few data issues. "
        "First, we noticed that some records in the 'NumAffected' column were blank or missing. "
        "We handled this by replacing the missing values with the median number of affected people for that specific type of incident. "
        "We also fixed some consistency issues with the DateTime column by converting it from a simple text string into a proper Datetime object. "
        "This was a very important step because it allowed us to extract the 'Month' and 'Year' for our trend analysis."
    )
    doc.add_paragraph(
        "We also derived a new metric called the 'Risk Score' to help the DRRMO prioritize their response. "
        "We created this by assigning a numerical weight to the Hazard Level (Critical = 4, High = 3, Medium = 2, Low = 1) "
        "and then multiplying it by the population of the barangay. "
        "This transformed our raw data into a strategic insight that ranks the most vulnerable areas in Biñan."
    )
    
    # --- 6. KPIs and metrics ---
    # Rubric Section 6
    doc.add_heading('6. KPIs and Metrics', level=1)
    doc.add_paragraph("The dashboard displays 10 key metrics to help the DRRMO monitor the city's performance:")
    
    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Table Grid'
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'KPI / Metric'
    hdr_cells3[1].text = 'Source Field'
    hdr_cells3[2].text = 'Update Frequency'
    hdr_cells3[3].text = 'Supports What Decision?'
    
    data3 = [
        ('Total Incidents', 'Incidents count', 'Real-time', 'Monitors current workload'),
        ('Total Affected', 'Sum(NumAffected)', 'Real-time', 'Checks the scale of impact'),
        ('Available Teams', 'Status column', 'Real-time', 'Identifies who can be deployed'),
        ('Total Capacity', 'Sum(Capacity)', 'Periodic', 'Checks total evacuation space'),
        ('Highest Risk Brgy', 'RiskScore', 'Periodic', 'Helps in long-term planning'),
        ('Evacuation Sites', 'Count(Facilities)', 'Periodic', 'Shows how many shelters are open'),
        ('Water Rescue Teams', 'Specialty', 'Real-time', 'Important for flood response'),
        ('Critical Incidents', 'Severity', 'Real-time', 'Highlights urgent emergencies'),
        ('Occupancy Rate', 'Occupants/Capacity', 'Real-time', 'Prevents overcrowding in sites'),
        ('Need Resupply', 'ResourcesAvailable', 'Real-time', 'Shows sites low on food/water')
    ]
    for row in data3:
        row_cells = table3.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        
    # --- 7. Dashboard design and explanation of visuals ---
    # Rubric Section 7
    doc.add_heading('7. Dashboard Design and Explanation of Visuals', level=1)
    doc.add_paragraph("We designed our charts to answer specific questions for the city responders:")
    
    doc.add_paragraph("1. Incident Trend Over Time: This line chart answers 'Are incidents increasing?'. It uses the Month and Incident count. It helps the city prepare for seasonal floods.")
    doc.add_paragraph("2. Incident Composition: A horizontal bar chart that answers 'Which hazard happens most often?'. It shows that floods are the biggest threat.")
    doc.add_paragraph("3. Risk Ranking: This bar chart uses the Risk Score to rank barangays. It tells the city where to focus their flood mitigation budget.")
    doc.add_paragraph("4. Gap Analysis: This chart compares affected people vs. available capacity. It helps the city decide if they need to open more evacuation sites.")
    doc.add_paragraph("5. Severity Pie Chart: This shows the ratio of Critical vs. Low incidents. It gives an immediate idea of the overall threat level.")
    doc.add_paragraph("6. Interactive Map: We used Folium to plot every incident and facility. We used custom icons for hospitals and schools so the user can easily find help.")
    
    # (Note: Data Entry is included as part of the overall design/functionality)
    doc.add_paragraph(
        "A key operational feature is the interactive data entry forms. "
        "The staff can log new incidents and add new facilities directly in the notebook. "
        "When they submit the form, the records are saved and the dashboard updates automatically."
    )

    # --- 8. Interaction and analytical navigation (OLAP) ---
    # Rubric Section 8
    doc.add_heading('8. Interaction and Analytical Navigation', level=1)
    doc.add_paragraph(
        "Our dashboard supports multidimensional analysis through interactive filters. "
        "The users can perform the following OLAP operations:"
    )
    doc.add_paragraph("Slice: Filtering the dashboard to show only the 'Flood' incidents.")
    doc.add_paragraph("Dice: Filtering to show 'Critical' Floods during a specific 'Month'.")
    doc.add_paragraph("Drill-down: Moving from the total KPI counts down to the exact street address on the interactive map or the details table.")
    doc.add_paragraph("Roll-up: Aggregating individual incident reports into a monthly trend view to see the big picture.")
    
    # --- 9. DRRM insights and recommendations ---
    # Rubric Section 9
    doc.add_heading('9. DRRM Insights and Recommendations', level=1)
    doc.add_paragraph(
        "From our analysis, we found that coastal barangays like Malaban and De La Paz are consistently at high risk. "
        "The gap analysis also shows that these areas often have overcrowded evacuation centers. "
        "Recommendation: The city should invest in bigger evacuation centers in these coastal areas. "
        "We also recommend placing specialized water rescue teams permanently in these barangays during the months of July to October."
    )
    
    # --- 10. Limitations and future improvements ---
    # Rubric Section 10
    doc.add_heading('10. Limitations and Future Improvements', level=1)
    doc.add_paragraph(
        "The current system uses simulated data that we made as realistic as possible for Biñan City. "
        "In the future, we can improve this by connecting it to a live weather database or an automated SMS alert system. "
        "We can also add more dimensions like weather humidity and rainfall data to make the predictions better."
    )

    doc.save('DRRMFinalReport_Funelas_Rivera_Torculas.docx')

if __name__ == "__main__":
    create_report()
